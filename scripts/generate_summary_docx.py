"""Generate the Rapaport family summary as Word documents (English + Polish)."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

OUT_DIR = Path('docs/research/deliverables')
OUT_DIR.mkdir(parents=True, exist_ok=True)


def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def add_title(doc, text, size=24):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x6B, 0x1F, 0x1F)
    return p


def add_h1(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(16)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x6B, 0x1F, 0x1F)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_h2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(13)
    r.font.bold = True
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_para(doc, text, italic=False, bold=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.italic = italic
    r.bold = bold
    p.paragraph_format.space_after = Pt(4)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        r = p.add_run(item)
        r.font.size = Pt(11)


def add_table(doc, headers, rows, col_widths_cm=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        r.font.bold = True
        r.font.size = Pt(10)
        set_cell_bg(hdr[i], 'E8D8C8')
    for i, row in enumerate(rows):
        cells = table.rows[i + 1].cells
        for j, val in enumerate(row):
            cells[j].text = ''
            p = cells[j].paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(9.5)
    if col_widths_cm:
        for row in table.rows:
            for i, w in enumerate(col_widths_cm):
                row.cells[i].width = Cm(w)


# ─────────────────────────────────────────────────────────────────────
# ENGLISH CONTENT
# ─────────────────────────────────────────────────────────────────────
EN = {
    "title": "Rapaport Family Research",
    "subtitle": "State of knowledge as of 1 June 2026",
    "intro": (
        "This document summarises what we know about the Rapaport family, how we know it, "
        "and which earlier conclusions we have had to retract. It is a working brief for "
        "Dalia, Dana, Daniel and Doron Rapaport — and for Basia, Magda, Kasia and Ms. Kasia, "
        "who have made the documented findings below possible. Drafted by Claude Code "
        "after the third WhatsApp-export ingest of the family-tree project. Every claim "
        "below is sourced; inferences are flagged."
    ),
    "h1_paternal": "1. The paternal line (from Moses Saul to us)",
    "paternal_intro": (
        "We can now place six confirmed generations of the paternal line. Until May 2026 "
        "only Berisz's name was known, with no documented parents; in the past six weeks "
        "we have added Moses Saul + Menukha as the apical couple (great-great-grandparents) "
        "and four siblings of Berisz, all corroborated by primary documents."
    ),
    "paternal_rows": [
        ("Great-great-grandparents",
         "Moses (Moshe) Saul Rapaport + Menukha",
         "Merchant family in Radomyśl Wielki. Moshe Rappaport served on the Radomyśl "
         "Wielki Jewish community Board of Trustees in 1897–1900 and 1901–1905.",
         "Yad Vashem PoT for Berisz (1953); three Tarnów marriage registers of their "
         "daughters (1919, 1920, 1923) — each explicitly names \"Moses Rapaport and "
         "Menucha\" as the bride's parents; Radomyśl Wielki Yizkor Book Part I (Board "
         "of Trustees listing)."),
        ("Great-grandparents",
         "Berisz / Bernard / Dov Rapaport (1886–1942) + Rebeka née Griffel (1888–1942)",
         "Berisz born 30 July 1886 in Radomyśl Wielki. Industrialist in Stanisławów "
         "from 1918. Married Rebeka civilly in Nadwórna on 21 December 1911. By 1939 "
         "the family lived in Przemyśl. Berisz was deported to Auschwitz and murdered "
         "there. Rebeka was killed in Przemyśl on a fabricated railway-bomb "
         "accusation, probably 1942.",
         "USHMM passport file RG-31.064M (Berisz's 1924 + 1927 Stanisławów applications); "
         "Yad Vashem PoTs 90394 + 90395 (1953); Auschwitz Museum victim record #188161, "
         "transport #689; Edward Gelles \"Griffel of Nadworna\" pedigree entry #8; "
         "Basia's research in the Nadwórna marriage register; Lusia's memoir page 24."),
        ("Grandparents",
         "David Mendel (Memek) Rapaport (1911–1990) + Leah Lusia née Weitzner (1913 or 1916 – 1996)",
         "David born 25 December 1911 in Nadwórna. Forestry engineer, polyglot. "
         "Lusia born in Bolechów. The couple married in Bolechów in 1935. Lusia ran "
         "the Willa \"Helin\" hotel in Muszyna pre-WWII; survived in Lwów as \"Maria "
         "Cizlik\". David escaped Galicia, reached Brussels April 1946, then sailed to "
         "Mandatory Palestine on the \"Theodor Herzl\" in April 1947.",
         "Nadwórna 1911 birth certificate; 1946 Brussels DIPIS card; 1935 Bolechów "
         "marriage certificate; CKŻP survivor card for Lusia (Katowice 1946); Lusia's "
         "memoir; the descendants tree."),
        ("Parents",
         "Dov (Bernard) Rapaport (b. 28 Aug 1946 Brussels) + Dalia née Goldfischer (b. 1952)",
         "Dov born in Brussels. Married Dalia in Haifa, 16 January 1974.",
         "Brussels DIPIS card; family records."),
        ("Children",
         "Dana (b. 1979 Haifa), Doron (b. 1981 Haifa), Daniel (b. 1983 Haifa)",
         "Living family.",
         "Family records."),
    ],
    "siblings_intro": (
        "Berisz had four documented sisters, all born in Radomyśl Wielki. Three of their "
        "marriages took place in Tarnów and each marriage register entry independently "
        "names Moses + Menukha as the bride's parents — the three-fold corroboration of "
        "the apical couple is the strongest evidence we have."
    ),
    "siblings_rows": [
        ("Alte Leja Rapaport (1882–1942)",
         "Born 28 June 1882. Married Turkel; the family lived in Vienna; seven "
         "documented children. Murdered at Auschwitz, 1942.",
         "Family-built descendants tree; cross-references with the Turkel-Tribe "
         "genealogy site (turkel.org.il)."),
        ("Jente Rapaport (b. 1887)",
         "Widowed by 1920 (first husband unknown). Remarried Mendel Eachem Horowitz "
         "of Sokal in Tarnów, 12 December 1920.",
         "Tarnów Jewish marriage register entry 98, 12 Dec 1920 — Archiwum Narodowe "
         "w Krakowie, Oddział w Tarnowie."),
        ("Freida Amalia Rapaport (1888–1942)",
         "First husband a Nussbaum, deceased by mid-1919 (their son Zwi Ayalon "
         "Nussbaum was born 7 May 1919 in Przemyśl). Remarried Markus Elias Kleinbaum "
         "in Tarnów, 8 June 1919. Died Przemyśl 1942.",
         "Tarnów Jewish marriage register entry 40, 8 June 1919; descendants tree."),
        ("Rebeka Rapaport the younger (b. 1896)",
         "Born Radomyśl Wielki. Married Sane Zylberfenig of Płońsk in Tarnów, "
         "4 September 1923. Later moved to Austria. NOT to be confused with "
         "Berisz's wife Rebeka née Griffel.",
         "Tarnów Jewish marriage register entry 101, 4 Sept 1923."),
    ],
    "cousin_note": (
        "Documented paternal first cousin who survived: Zwi (Hirsch Heschel) Ayalon "
        "Nussbaum — born 7 May 1919 Przemyśl, son of Freida Amalia, survived the "
        "Holocaust, emigrated to Israel, died 27 November 2001 in Haifa. His Israeli "
        "surname \"Ayalon\" is a Hebraisation of \"Nussbaum\"."
    ),
    "h1_maternal": "2. The maternal-side Griffel line",
    "maternal_text": (
        "Well documented via Edward Gelles's published pedigree \"Griffel of "
        "Nadworna\" (Balliol College Oxford archives). Apical: Dawid Griffel + Tauba → "
        "Eliezer \"Zeida\" Griffel (1850–1918) m. Sara Matel née Chajes (d. 1940) in "
        "Nadwórna 1892. Eliezer was head of the Nadwórna Jewish community (Av Kehillah) "
        "and an industrialist (timber, oil). Of his ten children, our line runs through "
        "Rivka (b. 1888) — our great-grandmother Rebeka, who married Berisz. The branch "
        "includes the Vaad ha-Hatzala Istanbul rescuer Dr Jacob Griffel, the famous "
        "memoirist Yehuda Nir, and the journalist Sarah Maslin Nir (NYT). All connections "
        "documented by Edward Gelles."
    ),
    "h1_holocaust": "3. The Holocaust — what is documented for each",
    "holocaust_rows": [
        ("Berisz / Bernard Dov Rapaport",
         "Murdered at Auschwitz",
         "Auschwitz Museum victim record #188161, transport #689"),
        ("Rebeka née Griffel",
         "Killed in Przemyśl, ~1942, on a fabricated railway-bomb accusation",
         "Lusia's memoir page 24"),
        ("Lota Rapaport (David's sister)",
         "Listed in the Kraków ghetto, 1940 (Reel 13/List 177/Entry 72, Ancestry "
         "\"Ten Ghettos\" DB). Per Lusia's memoir: later betrayed and arrested in "
         "Lwów with her husband; never returned.",
         "Ancestry/JewishGen \"Ten Jewish Ghettos 1939–1942\"; Lusia's memoir Chapter D"),
        ("Freida Amalia Rapaport-Kleinbaum",
         "Died Przemyśl 1942",
         "Descendants tree"),
        ("Alte Leja Rapaport-Turkel",
         "Murdered Auschwitz 1942",
         "Descendants tree"),
        ("Israel Menachem Turkel",
         "Killed Lwów 1942",
         "Descendants tree"),
        ("Siegfried Turkl",
         "Killed Belgium 12 January 1945",
         "Descendants tree"),
        ("Michael + Rachel Rosenfeld",
         "Killed Przemyśl 1944",
         "Descendants tree"),
    ],
    "h2_survivors": "Who survived",
    "survivors_rows": [
        ("David Memek Rapaport",
         "Escaped Galicia; reached Brussels April 1946; sailed to Palestine via the "
         "\"Theodor Herzl\" April 1947. The precise route through USSR / Iran / Italy "
         "is not fully reconstructed."),
        ("Lusia Rapaport",
         "Survived Lwów under the false Polish-Catholic identity \"Maria Cizlik\"."),
        ("Shimon Rapaport (David & Lusia's son, b. 22 June 1937 Lwów)",
         "Survived with Lusia. Became a senior reporter for the Israeli daily Ma'ariv, "
         "serving for many years as Northern-region correspondent — covering security "
         "events, industry, society and regional economics. Deceased (זכרונו לברכה)."),
        ("Zwi Ayalon Nussbaum (Berisz's nephew, b. 1919 Przemyśl)",
         "Survived; lived in Haifa, where he died 27 November 2001."),
    ],
    "h1_places": "4. Places in the archive — what each one is and why we hold it",
    "places_intro": (
        "Every place listed here is connected to the family by a sourced event. Places "
        "added during the May 2026 \"Ben-Zion philosopher\" hypothesis (Nowy Sącz, "
        "Żmigród, Bełżec, Gorlice, Kraków-as-Berisz-residence) were removed when the "
        "philosopher identification was retracted; they do not appear below."
    ),
    "places_rows": [
        ("Radomyśl Wielki, Poland",
         "Birthplace of Berisz Rapaport and his four sisters (Alte Leja 1882, Jente "
         "1887, Freida Amalia 1888, younger Rebeka 1896). Moses Saul Rapaport sat on "
         "the Jewish community board here 1897–1905.",
         "1924 passport application; three Tarnów marriage records (1919, 1920, 1923); "
         "Radomyśl Wielki Yizkor Book Part I."),
        ("Tarnów, Poland",
         "Where the three Tarnów Jewish marriage records were registered (1919 "
         "Freida × Kleinbaum, 1920 Jente × Horowitz, 1923 younger Rebeka × "
         "Zylberfenig). Two Moshe Saul Rapaport tombstones survive in the Tarnów "
         "Jewish cemetery — exact match with our line not yet established.",
         "Archiwum Narodowe w Krakowie, Oddział w Tarnowie."),
        ("Stanisławów (today Ivano-Frankivsk, Ukraine)",
         "Berisz's primary residence from 1918. Address: Kościuszki 4. Place of his "
         "1924 passport application + 1927 ID renewal. David attended the State First "
         "Gymnasium here in 1926 (his signature survives on the 150-year-of-US-"
         "independence dedication page).",
         "USHMM RG-31.064M; 1926 gimnazjum signature page."),
        ("Nadwórna (today Nadvirna, Ukraine)",
         "Birthplace of David Mendel (25 Dec 1911) and his mother Rebeka née Griffel "
         "(1888). Home of the Griffel and Chajes families. Eliezer Griffel's "
         "property on Śródmieście Street was auctioned in November 1938 (the 17-heir "
         "auction notice documents Rebeka as one of the heirs).",
         "Nadwórna 1911 birth certificate; 1938 Akcyjny Bank Hipoteczny auction notice."),
        ("Bolechów (today Bolekhiv, Ukraine)",
         "Birthplace of Lusia Weitzner and her siblings (Feige 1911, Lea 1913 or "
         "1916, Moses 1916). The Weitzner family lived in Bolechów Ruski. Place of "
         "the 1935 marriage of Lea Weitzner × Dawid Mendel Rapoport.",
         "Bolechów Jewish vital records via ŻIH Warsaw; 1935 marriage certificate."),
        ("Muszyna, Poland (\"Mosina\" in family memory)",
         "Spa resort where Lusia ran the Willa \"Helin\" hotel before WWII; Zygmunt "
         "Griffel (David's first cousin once removed) operated a sawmill here.",
         "Lusia's memoir; 1933 Muszyna Commune Council minutes; 1938 Przegląd Drzewny."),
        ("Przemyśl, Poland",
         "Where Berisz and Rebeka lived by 1939. When war broke out, Lusia and toddler "
         "Shimon went there to join them; David followed. Berisz was deported from "
         "Przemyśl to Auschwitz; Rebeka was killed in Przemyśl. Birthplace of cousin "
         "Zwi Ayalon Nussbaum (1919).",
         "Lusia's memoir; Yad Vashem PoT for Regina Rivka; Basia's notes (Bernard "
         "Rapaport in Przemyśl records ca. 1910)."),
        ("Lwów (today Lviv, Ukraine)",
         "David's residence as of 1 January 1938 (per Brussels DIPIS card). Where "
         "David and Lusia survived under false identities. Lusia lived at Legionów 24 "
         "(today Prospekt Svobody 24) as \"Maria Cizlik\". Lota was arrested and "
         "disappeared from here.",
         "Brussels DIPIS card; Lusia's memoir."),
        ("Olchowce (województwo lwowskie)",
         "Place of the religious marriage of Berisz × Rebeka, 1908, per the Gelles "
         "pedigree. Civil marriage followed at Nadwórna on 21 Dec 1911.",
         "Edward Gelles \"Griffel of Nadworna\"."),
        ("Sokal (today Sokalʹ, Ukraine)",
         "Hometown of Mendel Horowitz, second husband of Berisz's sister Jente.",
         "1920 Tarnów marriage register entry 98."),
        ("Płońsk, central Poland",
         "Hometown of Sane Zylberfenig, husband of Berisz's youngest sister Rebeka.",
         "1923 Tarnów marriage register entry 101."),
        ("Brussels, Belgium",
         "Where David ended his escape journey on 9 April 1946 (date on the DIPIS "
         "card). Dov was born here 28 August 1946.",
         "DIPIS card."),
        ("Auschwitz / Oświęcim, Poland",
         "Where Berisz was murdered (victim #188161). His sister Alte Leja was also "
         "murdered there in 1942.",
         "Auschwitz Museum victim database."),
        ("Sète, France & Marseille, France",
         "Embarkation port (Sète) and shipyard (Marseille) of the \"Theodor Herzl\" "
         "— the ship that brought David, Lusia, Shimon and infant Dov from Brussels "
         "to Mandatory Palestine in April 1947.",
         "Brussels DIPIS card + family memory + Mossad LeAliyah Bet historical "
         "records."),
        ("Cyprus (Karaolos detention camp)",
         "British detention camp where Theodor Herzl passengers including the "
         "Rapaports were held before being released to Palestine.",
         "Mossad LeAliyah Bet records."),
        ("Atlit, Mandatory Palestine",
         "British landing-and-detention camp on arrival in Palestine.",
         "Mossad LeAliyah Bet records."),
        ("Haifa, Israel",
         "Where David, Lusia, Shimon, Dov settled. David died here 29 August 1990. "
         "Lusia died here 1996. Dana, Doron, Daniel all born here.",
         "Family records."),
        ("Vienna, Austria",
         "Where Berisz's sister Alte Leja settled with her husband Turkel; their "
         "Turkel descendants were born here (Auersperggasse 9 + Rembrandtstrasse 3). "
         "Berisz travelled here on business in 1924.",
         "1924 passport application; descendants tree."),
    ],
    "h1_documents": "5. Documents we hold",
    "documents_text": (
        "39 catalogued documents in the live archive at rapaportfamily.github.io. "
        "Highlights:"
    ),
    "documents_bullets": [
        "Berisz's 1924 + 1927 passport applications, Stanisławów (USHMM RG-31.064M, frames 375-380)",
        "Yad Vashem Pages of Testimony filed 1953: Bernard Dov Rapaport (file 90394) and Regina Rivka Rapaport (file 90395)",
        "Auschwitz Museum victim record #188161, transport #689 — Berisz",
        "Ancestry / JewishGen \"Ten Jewish Ghettos 1939–1942\" — Lotte Rapaport, Kraków ghetto 1940 (Reel 13, List 177, Entry 72)",
        "Tarnów Jewish marriage register pages — entries 40 (1919), 98 (1920), 101 (1923) — the three sister marriages",
        "Nadwórna 1911 birth certificate of David Mendel Rapaport",
        "Nadwórna 1888 birth certificate of Rebeka Griffel",
        "Bolechów 1935 marriage certificate (Lea Weitzner × Dawid Mendel Rapoport)",
        "Bolechów 1932 marriage certificate (Feige Weitzner × Israel Englander)",
        "Stanisławów 1926 gymnasium signature page — David's own handwriting (Klasa IV)",
        "1938 Akcyjny Bank Hipoteczny auction notice listing all 17 heirs of Eliezer Griffel",
        "1934-35 Sejm Album feature on Zygmunt Griffel's Lwów timber business",
        "1938 Przegląd Drzewny article on Zygmunt Griffel's Muszyna sawmill",
        "CKŻP survivor registration cards: Lusia (Nr 151738) + Shimon (Nr 151337), Katowice 1 July 1946",
        "Brussels DIPIS card — David Rapaport, 9 April 1946",
        "Lusia's 90-page Hebrew memoir \"Lusia's Story\", with Polish and English translations",
        "Edward Gelles, \"Griffel of Nadworna\" pedigree + \"Facets of My Family History\" Parts 1 & 2 (Balliol College Oxford)",
        "Family-built descendants trees of Moses Saul Rapaport and of Dawid Griffel (Basia, May 2026)",
        "Ma'ariv newspaper archive entry, 30 August 1990 — referenced under שמעון רפפורט one day after David's death (National Library of Israel)",
    ],
    "h1_retractions": "6. What we thought was right and was wrong",
    "retractions_intro": (
        "We have had to retract several conclusions. Each retraction protected us from "
        "building on false ground; the lessons are part of the value of the project."
    ),
    "retractions_rows": [
        ("Berisz born in Tarnów 6 August 1884 as \"Benzion\", son of Mojżesz Saul + Rywka Schiff",
         "Berisz's own 1924 passport application: born 30 July 1886 in Radomyśl Wielki, "
         "father Moses. The Tarnów Benzion was a different Rapaport family."),
        ("Berisz = the published Hebrew philosopher Ben-Zion Rappaport, author of "
         "\"Nature and Spirit\" (Mossad Bialik 1953); deported from Nowy Sącz to Bełżec; "
         "first wife from Żmigród; son Moshe Hacohen; sister Sarah Mahler",
         "Lusia's memoir places Berisz in Przemyśl, not Sącz/Bełżec, and never mentions "
         "philosophy or books. The Tarnów-Yizkor Ben-Zion was a different person who "
         "shared only birth city + year. Auschwitz Museum confirms Berisz died at "
         "Auschwitz."),
        ("Berisz died in Nadwórna at the 6 October 1941 Bukowinka Forest mass aktion "
         "or in the 24 October 1942 ghetto liquidation",
         "Auschwitz Museum victim record #188161 — Berisz was deported to Auschwitz "
         "and murdered there."),
        ("Rebeka Griffel was killed at the Bukowinka Forest aktion 1941",
         "Lusia's memoir and Yad Vashem PoT both place her in Przemyśl during the war; "
         "she was killed there."),
        ("The 1941 \"Lea Rapaport divorce\" document concerned our Lusia",
         "Reading the document directly: this is a different Lea — Lea bat Shimon "
         "HaKohen Rapaport × Michael Sigal in Petach Tikvah, Mandatory Palestine, "
         "certificate No. H 13898. Our Lusia was in Lwów hiding as Maria Cizlik in 1941."),
    ],
    "h1_questions": "7. Open questions, ranked by leverage",
    "questions_rows": [
        ("Pull the full Ancestry record for Lota in the Kraków ghetto",
         "Reel 13 / List 177 / Entry 72 in the \"Ten Ghettos\" database may include "
         "her parents' names, husband's name, address, and occupation — likely enough "
         "to finally identify her husband."),
        ("Resolve the Moses Saul Rapaport tombstone",
         "Two Moshe Saul Rapaport tombstones in Tarnów cemetery: one died 11 Aug "
         "1933, the other 30 Oct 1931 from Dąbrowa, 7th-generation descendant of "
         "Shabtai HaKohen (\"the Shach\"). Either could be ours — resolving it ties "
         "our line to a specific rabbinic descent."),
        ("Trace Berisz's pre-Stanisławów residence",
         "Berisz lived in Stanisławów \"since 1918\". Where did the family live "
         "1886–1918? Radomyśl Wielki for some of those years (children born there "
         "until 1896); the gap 1896–1918 is undocumented."),
        ("Identify the first husbands of Freida Amalia (Nussbaum, ca. 1917-18) and "
         "Jente (unknown, before 1920)",
         "Tarnów or Przemyśl marriage records pre-1919 should help."),
    ],
    "h1_status": "8. Current state of the live archive",
    "status_text": (
        "Hosted at rapaportfamily.github.io. Magic-link auth for family; guest URL "
        "for sharing. As of 1 June 2026: 113 people, 44 places, 39 documents, 541 "
        "chat messages, Lusia's memoir as flip-book in Hebrew, English and Polish. "
        "Installed as a Progressive Web App on phones and laptops, with automatic "
        "update detection so new findings reach the family without manual refresh."
    ),
    "thanks": (
        "With gratitude to Basia (genealogical research in Polish archives), Magda "
        "(who connected us), Kasia (Jewish Historical Institute, Warsaw), Ms. Kasia "
        "(French translation of the DIPIS card), Edward Gelles (the published "
        "Griffel pedigree), and the late Lusia Rapaport, whose memoir keeps the "
        "family's story alive."
    ),
}


def build_doc(content, out_path):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    add_title(doc, content["title"], size=28)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(content["subtitle"])
    r.italic = True
    r.font.size = Pt(13)

    doc.add_paragraph()
    add_para(doc, content["intro"], italic=True)

    # 1. Paternal line
    add_h1(doc, content["h1_paternal"])
    add_para(doc, content["paternal_intro"])
    add_table(doc,
              ["Generation", "Names", "Verified facts", "How we know"],
              content["paternal_rows"],
              col_widths_cm=[3.0, 4.5, 5.0, 4.0])

    add_h2(doc, "Berisz's siblings (children of Moses Saul + Menukha)")
    add_para(doc, content["siblings_intro"])
    add_table(doc,
              ["Sibling", "Verified facts", "Source"],
              content["siblings_rows"],
              col_widths_cm=[4.5, 7.0, 5.0])

    add_para(doc, content["cousin_note"], italic=True)

    # 2. Maternal line
    add_h1(doc, content["h1_maternal"])
    add_para(doc, content["maternal_text"])

    # 3. Holocaust
    add_h1(doc, content["h1_holocaust"])
    add_table(doc,
              ["Person", "Fate", "Source"],
              content["holocaust_rows"],
              col_widths_cm=[4.5, 7.0, 5.0])

    add_h2(doc, content["h2_survivors"])
    add_table(doc,
              ["Person", "How / where"],
              content["survivors_rows"],
              col_widths_cm=[5.5, 11.0])

    # 4. Places
    add_h1(doc, content["h1_places"])
    add_para(doc, content["places_intro"])
    add_table(doc,
              ["Place", "Why this place is in the archive", "Documentary source"],
              content["places_rows"],
              col_widths_cm=[4.0, 7.0, 5.5])

    # 5. Documents
    add_h1(doc, content["h1_documents"])
    add_para(doc, content["documents_text"])
    add_bullets(doc, content["documents_bullets"])

    # 6. Retractions
    add_h1(doc, content["h1_retractions"])
    add_para(doc, content["retractions_intro"])
    add_table(doc,
              ["Claim that was wrong", "How it was refuted"],
              content["retractions_rows"],
              col_widths_cm=[8.0, 8.5])

    # 7. Open questions
    add_h1(doc, content["h1_questions"])
    add_table(doc,
              ["Next step", "Why it matters"],
              content["questions_rows"],
              col_widths_cm=[6.0, 10.5])

    # 8. Status
    add_h1(doc, content["h1_status"])
    add_para(doc, content["status_text"])

    doc.add_paragraph()
    add_para(doc, content["thanks"], italic=True)

    doc.save(out_path)
    return out_path


# Build English
en_path = OUT_DIR / "Rapaport_Family_Summary_EN.docx"
build_doc(EN, en_path)
print(f"wrote: {en_path}")
print("EN done. Polish in next step.")
