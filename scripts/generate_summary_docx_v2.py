"""V2: Word docs with Executive Summary first, then full content. EN + PL.

Structure:
  Cover (title + date + author)
  Executive Summary (1 page)
  ---page break---
  Full detail (the long version)

Verifies output structure before declaring done.
"""
import sys
sys.path.insert(0, 'scripts')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

OUT_DIR = Path('docs/research/deliverables')
OUT_DIR.mkdir(parents=True, exist_ok=True)


def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def add_title(doc, text, size=28, color=(0x6B, 0x1F, 0x1F)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = True
    r.font.color.rgb = RGBColor(*color)
    return p


def add_subtitle(doc, text, size=13):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(size)
    return p


def add_h1(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(16)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x6B, 0x1F, 0x1F)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_h2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(13)
    r.font.bold = True
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_para(doc, text, italic=False, bold=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.italic = italic
    r.bold = bold
    p.paragraph_format.space_after = Pt(4)
    return p


def add_bullets(doc, items, bold_first_phrase=False):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        if bold_first_phrase and ' — ' in item:
            left, right = item.split(' — ', 1)
            r1 = p.add_run(left)
            r1.font.size = Pt(11)
            r1.font.bold = True
            r2 = p.add_run(' — ' + right)
            r2.font.size = Pt(11)
        else:
            r = p.add_run(item)
            r.font.size = Pt(11)


def add_table(doc, headers, rows, col_widths_cm=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        r.font.bold = True
        r.font.size = Pt(10)
        set_cell_bg(hdr[i], 'E8D8C8')
    for i, row in enumerate(rows):
        cells = table.rows[i + 1].cells
        for j, val in enumerate(row):
            cells[j].text = ''
            p = cells[j].paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(9.5)
    if col_widths_cm:
        for row in table.rows:
            for i, w in enumerate(col_widths_cm):
                row.cells[i].width = Cm(w)


def page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


# ─── Re-use the long content from V1 (EN + PL dicts) ───
from generate_summary_docx import EN
from generate_summary_docx_pl import PL


# ─── Executive Summary content ───
EN_TLDR = {
    "headline": "Six generations of the paternal line confirmed. Berisz Rapaport was born 30 July 1886 in Radomyśl Wielki, not Tarnów; murdered at Auschwitz. Rebeka Griffel killed in Przemyśl ~1942. Lota documented in Kraków ghetto 1940. The apical couple Moses Saul + Menukha is confirmed by THREE independent Tarnów marriage records of Berisz's sisters.",
    "key_findings_title": "Five things we now know that we didn't six weeks ago",
    "key_findings": [
        "**Berisz's exact birth** — 30 July 1886 in Radomyśl Wielki (USHMM passport file, found by Basia 25 May 2026). Earlier Tarnów-1884 reading was a different family.",
        "**The apical couple** Moses Saul Rapaport + Menukha — corroborated by THREE Tarnów marriage records of Berisz's sisters (1919, 1920, 1923), each naming \"Moses + Menucha\" independently.",
        "**Four sisters of Berisz** plus a surviving paternal first cousin (Zwi Ayalon Nussbaum, b. 1919 Przemyśl, d. 2001 Haifa). The Vienna Turkel branch via Alte Leja's marriage.",
        "**Lota in Kraków ghetto 1940** — Ancestry/JewishGen \"Ten Ghettos\" record, Reel 13/List 177/Entry 72. The first Holocaust-era documentation of David's sister.",
        "**Holocaust fates corrected**: Berisz at Auschwitz (Museum #188161, transport #689), not the Bukowinka mass aktion. Rebeka killed in Przemyśl per Lusia's memoir, not Nadwórna.",
    ],
    "retractions_title": "Five things we thought were true and weren't",
    "retractions": [
        "Berisz ≠ the Tarnów-1884 Benzion (different family).",
        "Berisz ≠ the published Hebrew philosopher Ben-Zion Rappaport. Different person — coincidence on city + birth year.",
        "Berisz didn't die in Nadwórna or Bełżec. He was deported to Auschwitz.",
        "Rebeka didn't die at the Bukowinka Forest aktion. She was killed in Przemyśl.",
        "The 1941 Lea Rapaport divorce paper isn't our Lusia — it's a different Lea Rapaport in Petach Tikvah.",
    ],
    "archive_status": "Live archive at https://rapaportfamily.github.io — 113 people · 44 places · 39 documents · 541 chat messages · Lusia's memoir as a flipbook in Hebrew, English, Polish · installable as a PWA with automatic updates.",
}

PL_TLDR = {
    "headline": "Potwierdziliśmy sześć pokoleń linii ojcowskiej. Berisz Rapaport urodził się 30 lipca 1886 r. w Radomyślu Wielkim, nie w Tarnowie; zamordowany w Auschwitz. Rebeka Griffel zabita w Przemyślu ok. 1942 r. Lota udokumentowana w getcie krakowskim 1940 r. Para protoplastów Mojżesz Saul + Menucha jest potwierdzona przez TRZY niezależne akty małżeństw tarnowskich sióstr Berisza.",
    "key_findings_title": "Pięć rzeczy, których nie wiedzieliśmy sześć tygodni temu",
    "key_findings": [
        "**Dokładna data urodzenia Berisza** — 30 lipca 1886 r. w Radomyślu Wielkim (akta paszportowe USHMM, znalezione przez Basię 25 maja 2026 r.). Wcześniejszy odczyt Tarnów-1884 dotyczył innej rodziny.",
        "**Para protoplastów** Mojżesz Saul Rapaport + Menucha — potwierdzona przez TRZY tarnowskie akty małżeństw sióstr Berisza (1919, 1920, 1923), z których każdy niezależnie wymienia „Moses + Menucha\".",
        "**Cztery siostry Berisza** plus ocalały kuzyn w pierwszej linii ojcowskiej (Zwi Ayalon Nussbaum, ur. 1919 Przemyśl, zm. 2001 Hajfa). Wiedeńska gałąź Turklów przez małżeństwo Alte Leji.",
        "**Lota w getcie krakowskim 1940 r.** — rekord Ancestry/JewishGen „Ten Ghettos\", Rolka 13/Lista 177/Pozycja 72. Pierwsza udokumentowana z czasu Holocaustu wzmianka o siostrze Dawida.",
        "**Skorygowane losy w Holokauście**: Berisz w Auschwitz (Muzeum #188161, transport #689), nie w masowej akcji w Bukowince. Rebeka zabita w Przemyślu według wspomnień Lusi, nie w Nadwórnej.",
    ],
    "retractions_title": "Pięć rzeczy, które uważaliśmy za prawdę, a okazały się błędne",
    "retractions": [
        "Berisz ≠ Benzion z Tarnowa 1884 (inna rodzina).",
        "Berisz ≠ opublikowany filozof hebrajski Ben-Zion Rappaport. Inna osoba — przypadkowa zbieżność miasta i roku urodzenia.",
        "Berisz nie zginął w Nadwórnej ani w Bełżcu. Został deportowany do Auschwitz.",
        "Rebeka nie zginęła w akcji w lesie Bukowinka. Została zabita w Przemyślu.",
        "Akt rozwodu Lei Rapaport z 1941 r. nie dotyczy naszej Lusi — to inna Lea Rapaport w Petach Tikwie.",
    ],
    "archive_status": "Archiwum na żywo pod adresem https://rapaportfamily.github.io — 113 osób · 44 miejsc · 39 dokumentów · 541 wiadomości czatu · wspomnienia Lusi jako interaktywna książka w hebrajskim, angielskim i polskim · instalowalne jako PWA z automatycznymi aktualizacjami.",
}


def render_bold_segments(p, text):
    """Render **bold** segments in a paragraph."""
    pieces = []
    rest = text
    while True:
        i = rest.find('**')
        if i == -1:
            pieces.append((rest, False))
            break
        if i > 0:
            pieces.append((rest[:i], False))
        rest = rest[i+2:]
        j = rest.find('**')
        if j == -1:
            pieces.append((rest, True))
            break
        pieces.append((rest[:j], True))
        rest = rest[j+2:]
    for text_part, bold in pieces:
        if not text_part:
            continue
        r = p.add_run(text_part)
        r.font.size = Pt(11)
        r.bold = bold


def add_bullet_with_bold(doc, item):
    p = doc.add_paragraph(style='List Bullet')
    render_bold_segments(p, item)


def build_exec_summary(doc, tldr, lang):
    if lang == 'EN':
        add_h1(doc, "Executive Summary")
    else:
        add_h1(doc, "Streszczenie wykonawcze")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.space_before = Pt(4)
    r = p.add_run(tldr["headline"])
    r.font.size = Pt(12)
    r.font.italic = True

    add_h2(doc, tldr["key_findings_title"])
    for finding in tldr["key_findings"]:
        add_bullet_with_bold(doc, finding)

    add_h2(doc, tldr["retractions_title"])
    for retr in tldr["retractions"]:
        p = doc.add_paragraph(style='List Bullet')
        r = p.add_run(retr)
        r.font.size = Pt(11)

    add_h2(doc, "Archive" if lang == 'EN' else "Archiwum")
    add_para(doc, tldr["archive_status"])


def build_full_body(doc, content, lang):
    if lang == 'EN':
        add_h1(doc, "Detailed report — full content follows")
    else:
        add_h1(doc, "Raport szczegółowy — pełna treść poniżej")

    add_para(doc, content["intro"], italic=True)

    # 1. Paternal
    add_h1(doc, content["h1_paternal"])
    add_para(doc, content["paternal_intro"])
    add_table(doc,
              (["Generation", "Names", "Verified facts", "How we know"] if lang == 'EN'
               else ["Pokolenie", "Imiona", "Potwierdzone fakty", "Skąd to wiemy"]),
              content["paternal_rows"],
              col_widths_cm=[3.0, 4.5, 5.0, 4.0])

    add_h2(doc, ("Berisz's siblings (children of Moses Saul + Menukha)" if lang == 'EN'
                 else "Rodzeństwo Berisza (dzieci Mojżesza Saula + Menuchy)"))
    add_para(doc, content["siblings_intro"])
    add_table(doc,
              (["Sibling", "Verified facts", "Source"] if lang == 'EN'
               else ["Rodzeństwo", "Potwierdzone fakty", "Źródło"]),
              content["siblings_rows"],
              col_widths_cm=[4.5, 7.0, 5.0])

    add_para(doc, content["cousin_note"], italic=True)

    # 2. Maternal
    add_h1(doc, content["h1_maternal"])
    add_para(doc, content["maternal_text"])

    # 3. Holocaust
    add_h1(doc, content["h1_holocaust"])
    add_table(doc,
              (["Person", "Fate", "Source"] if lang == 'EN'
               else ["Osoba", "Los", "Źródło"]),
              content["holocaust_rows"],
              col_widths_cm=[4.5, 7.0, 5.0])

    add_h2(doc, content["h2_survivors"])
    add_table(doc,
              (["Person", "How / where"] if lang == 'EN'
               else ["Osoba", "Jak / gdzie"]),
              content["survivors_rows"],
              col_widths_cm=[5.5, 11.0])

    # 4. Places
    add_h1(doc, content["h1_places"])
    add_para(doc, content["places_intro"])
    add_table(doc,
              (["Place", "Why this place is in the archive", "Documentary source"] if lang == 'EN'
               else ["Miejsce", "Dlaczego jest w archiwum", "Źródło"]),
              content["places_rows"],
              col_widths_cm=[4.0, 7.0, 5.5])

    # 5. Documents
    add_h1(doc, content["h1_documents"])
    add_para(doc, content["documents_text"])
    add_bullets(doc, content["documents_bullets"])

    # 6. Retractions
    add_h1(doc, content["h1_retractions"])
    add_para(doc, content["retractions_intro"])
    add_table(doc,
              (["Claim that was wrong", "How it was refuted"] if lang == 'EN'
               else ["Twierdzenie błędne", "Co je obaliło"]),
              content["retractions_rows"],
              col_widths_cm=[8.0, 8.5])

    # 7. Open questions
    add_h1(doc, content["h1_questions"])
    add_table(doc,
              (["Next step", "Why it matters"] if lang == 'EN'
               else ["Następny krok", "Dlaczego to ważne"]),
              content["questions_rows"],
              col_widths_cm=[6.0, 10.5])

    # 8. Status
    add_h1(doc, content["h1_status"])
    add_para(doc, content["status_text"])

    doc.add_paragraph()
    add_para(doc, content["thanks"], italic=True)


def build(content, tldr, lang, out_path):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Cover
    for _ in range(5):
        doc.add_paragraph()
    add_title(doc, content["title"], size=32)
    add_subtitle(doc, content["subtitle"])
    for _ in range(2):
        doc.add_paragraph()
    add_para(doc, ("Compiled by Claude Code for the Rapaport family research project."
                   if lang == 'EN'
                   else "Opracowane przez Claude Code dla projektu badawczego rodziny Rapaport."),
             italic=True)
    add_para(doc, ("1 June 2026" if lang == 'EN' else "1 czerwca 2026 r."), italic=True)

    page_break(doc)

    # Executive Summary
    build_exec_summary(doc, tldr, lang)

    page_break(doc)

    # Full body
    build_full_body(doc, content, lang)

    doc.save(out_path)
    return out_path


def verify(path, expected_lang):
    """Check the file opens and has expected structure."""
    d = Document(path)
    paras = [p for p in d.paragraphs if p.text.strip()]
    tables = d.tables
    table_rows = sum(len(t.rows) for t in tables)
    has_title = any(p.text.startswith(("Rapaport Family", "Badania nad rodziną")) for p in paras)
    has_exec = any(p.text.startswith(("Executive Summary", "Streszczenie wykonawcze")) for p in paras)
    has_detail = any(p.text.startswith(("Detailed report", "Raport szczegółowy")) for p in paras)
    has_8_sections = sum(1 for p in paras if p.text.startswith(("1.","2.","3.","4.","5.","6.","7.","8."))) >= 8
    has_thanks = any(("With gratitude" in p.text) or ("Z wdzięcznością" in p.text) for p in paras)
    return {
        "path": str(path),
        "size_bytes": path.stat().st_size,
        "paragraphs": len(paras),
        "tables": len(tables),
        "table_rows": table_rows,
        "has_title": has_title,
        "has_exec_summary": has_exec,
        "has_detail_section": has_detail,
        "has_8_sections": has_8_sections,
        "has_thanks": has_thanks,
        "lang_ok": has_title and has_exec and has_detail and has_8_sections and has_thanks,
    }


en_path = OUT_DIR / "Rapaport_Family_Summary_EN_v2.docx"
pl_path = OUT_DIR / "Rapaport_Family_Summary_PL_v2.docx"
build(EN, EN_TLDR, 'EN', en_path)
build(PL, PL_TLDR, 'PL', pl_path)

import json
print(json.dumps(verify(en_path, 'EN'), indent=2, ensure_ascii=False))
print(json.dumps(verify(pl_path, 'PL'), indent=2, ensure_ascii=False))
